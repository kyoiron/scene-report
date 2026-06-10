import streamlit as st
import google.generativeai as genai
import io
import os
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

st.set_page_config(page_title="現場照片勘驗生成器", page_icon="🔍", layout="wide")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Noto+Sans+TC:wght@400;500;700&display=swap');
html, body, [class*="css"] { font-family: 'Noto Sans TC', sans-serif; }
.stApp { background: #0f1117; color: #e0e0e0; }
.top-bar {
    background: linear-gradient(90deg,#1a1f35,#1e2640);
    border-bottom:1px solid #2a3150;
    padding:14px 32px; display:flex; align-items:center; gap:14px;
    margin:-1rem -1rem 2rem -1rem;
}
.top-bar h1{font-size:20px;font-weight:700;color:#f0f4ff;margin:0;}
.top-bar p{font-size:12px;color:#6b7db3;margin:0;}
.badge{display:inline-block;background:#1e3a70;color:#60a5fa;border-radius:5px;padding:2px 10px;font-size:12px;margin-bottom:8px;}
.field-label{font-size:13px;color:#8899bb;margin-bottom:4px;font-weight:500;}
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="top-bar">
  <div style="font-size:26px">🔍</div>
  <div>
    <h1>現場照片勘驗生成器</h1>
    <p>AI 輔助司法現場勘驗文字產製系統 ｜ Gemini Vision API</p>
  </div>
</div>
""", unsafe_allow_html=True)

# ── Session State 初始化 ───────────────────────────────────
if "photos" not in st.session_state:
    st.session_state.photos = []
if "api_key" not in st.session_state:
    st.session_state.api_key = os.environ.get("GEMINI_API_KEY", "")

# ── Sidebar ────────────────────────────────────────────────
with st.sidebar:
    st.markdown("### ⚙️ 系統設定")
    key_in = st.text_input("Gemini API Key", value=st.session_state.api_key,
                           type="password", placeholder="AIza...")
    if key_in:
        st.session_state.api_key = key_in
    if st.session_state.api_key:
        st.success("✓ API Key 已設定")
    else:
        st.warning("請輸入 Gemini API Key")
    st.markdown("---")
    st.markdown("""**使用說明**
1. 輸入 Gemini API Key
2. 填入案號
3. 上傳現場照片
4. 點擊「AI 生成」
5. 確認文字後匯出 docx""")
    st.markdown("---")
    st.caption("臺灣高雄地方檢察署\n資訊室 Demo 系統")

# ── Gemini 單張分析（含全案脈絡）────────────────────────
def call_gemini(api_key, image_bytes, mime_type, extra="",
                all_photos=None, target_index=0, case_num=""):
    """
    分析單張照片，但同時傳入所有照片供 AI 理解全案脈絡。
    all_photos: list of {bytes, mime, name}，為所有照片
    target_index: 本次要產生勘驗文字的照片編號（0-based）
    """
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-2.5-flash-lite")
    extra_text = f"\n\n補充指示：{extra}" if extra.strip() else ""
    case_hint = f"案號：{case_num}。" if case_num and case_num.strip() else ""

    # 有多張照片時，加入全案脈絡說明
    if all_photos and len(all_photos) > 1:
        total = len(all_photos)
        context_prompt = f"""你是一位資深司法現場勘驗人員。
{case_hint}本案共有 {total} 張現場照片，均為同一車禍事故現場。
以下會先提供所有照片供你理解整體事故全貌，
請針對【第 {target_index + 1} 張照片（{all_photos[target_index]['name']}）】撰寫勘驗描述。

要求：
1. 詳細描述該張照片的現場狀況、受損部位、損壞程度
2. 如有車輛，描述各部位受損情形及估計修復費用範圍
3. 如有跡證（血跡、碎片、標記物），詳細描述位置
4. 語氣符合司法文書規範，使用被動語態，語句簡潔精準
5. 分段描述，每段聚焦一個重點，不使用條列符號
6. 適當處可與其他角度照片相互呼應（如：「與編號X照片所示相符」）
7. 僅輸出該張照片的勘驗文字，不需標題或編號{extra_text}"""

        content = [context_prompt]
        for i, p in enumerate(all_photos):
            marker = f"【{'→ 本張目標' if i == target_index else f'參考照片 {i+1}'}：{p['name']}】"
            content.append(marker)
            content.append({"mime_type": p["mime"], "data": p["bytes"]})
    else:
        # 單張模式（向下相容）
        context_prompt = f"""你是一位資深司法現場勘驗人員，請根據此照片，以繁體中文撰寫正式勘驗描述。

要求：
1. 客觀描述現場狀況、受損部位、損壞程度
2. 如有車輛，描述各部位受損情形及估計修復費用範圍
3. 如有跡證（血跡、碎片、標記物），詳細描述位置
4. 語氣符合司法文書規範，使用被動語態，語句簡潔精準
5. 分段描述，每段聚焦一個重點，不使用條列符號{extra_text}"""
        content = [context_prompt, {"mime_type": mime_type, "data": image_bytes}]

    response = model.generate_content(content)
    return response.text

# ── Gemini 全案關聯分析：所有照片一次送入 ─────────────────
def call_gemini_all(api_key, photos, case_num=""):
    import json
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel("gemini-2.5-flash-lite")
    total = len(photos)
    case_hint = f"案號：{case_num}。" if case_num.strip() else ""
    system_prompt = f"""你是一位資深司法現場勘驗人員。
{case_hint}以下共 {total} 張照片，均為同一車禍事故現場的不同角度或部位。

請先綜合理解整個事故現場全貌，再針對每張照片分別撰寫勘驗描述。
各照片之勘驗文字應具有關聯性，避免重複描述，並在適當處相互呼應。

輸出格式（必須嚴格遵守，僅輸出 JSON 不要有其他文字）：
{{
  "overview": "整體事故現場概述（2-3句）",
  "photos": [
    {{"index": 1, "text": "第1張照片勘驗文字"}},
    {{"index": 2, "text": "第2張照片勘驗文字"}}
  ]
}}

每張照片勘驗文字需符合司法文書規範，使用被動語態，分段描述，不使用條列符號。"""

    content_list = [system_prompt]
    for i, p in enumerate(photos):
        content_list.append(f"\n【照片編號 {i+1}：{p['name']}】")
        content_list.append({"mime_type": p["mime"], "data": p["bytes"]})

    response = model.generate_content(content_list)
    raw = response.text.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    raw = raw.strip()

    data = json.loads(raw)
    results = [""] * total
    overview = data.get("overview", "")
    for item in data.get("photos", []):
        idx = item.get("index", 0) - 1
        if 0 <= idx < total:
            text = item.get("text", "")
            if idx == 0 and overview:
                results[idx] = f"【現場概述】{overview}\n\n{text}"
            else:
                results[idx] = text
    return results


# ── 工具函式：清除所有 widget state ──────────────────────
def clear_widget_states():
    for k in list(st.session_state.keys()):
        if k.startswith("txt_") or k.startswith("prompt_"):
            del st.session_state[k]

# ── 案號 + 上傳 ───────────────────────────────────────────
st.markdown('<div class="field-label">📋 案號</div>', unsafe_allow_html=True)
st.text_input("", placeholder="例：115重訴135",
              label_visibility="collapsed", key="case_num")

st.markdown('<div class="field-label">📷 上傳現場照片（支援多張）</div>', unsafe_allow_html=True)
uploaded = st.file_uploader("", type=["jpg","jpeg","png","webp"],
                             accept_multiple_files=True,
                             label_visibility="collapsed", key="uploader")
if uploaded:
    existing = {p["name"] for p in st.session_state.photos}
    for f in uploaded:
        if f.name not in existing:
            st.session_state.photos.append({
                "name": f.name, "bytes": f.read(),
                "mime": f.type or "image/jpeg",
                "text": "", "prompt": ""
            })
            existing.add(f.name)

if st.session_state.photos:
    st.markdown('<div class="field-label">已上傳照片</div>', unsafe_allow_html=True)
    thumb_cols = st.columns(min(len(st.session_state.photos), 6))
    for i, p in enumerate(st.session_state.photos):
        with thumb_cols[i % 6]:
            st.image(Image.open(io.BytesIO(p["bytes"])), width=80, caption=f"#{i+1}")
    if st.button("🗑 清除所有照片"):
        st.session_state.photos = []
        clear_widget_states()
        st.rerun()

st.markdown("---")

# ── 照片卡片區 ────────────────────────────────────────────
if not st.session_state.photos:
    st.markdown("""<div style="text-align:center;padding:40px;color:#3a4a6b;">
        <div style="font-size:40px;margin-bottom:8px">📂</div>
        <div>請先上傳現場照片，系統將自動產製勘驗文字</div>
    </div>""", unsafe_allow_html=True)
else:
    for i in range(len(st.session_state.photos)):
        p = st.session_state.photos[i]
        st.markdown(f'<div class="badge">編號 {i+1}｜{p["name"]}</div>',
                    unsafe_allow_html=True)
        c1, c2 = st.columns([1, 2])

        with c1:
            st.image(Image.open(io.BytesIO(p["bytes"])), use_container_width=True)
            o1, o2, o3 = st.columns(3)
            with o1:
                if st.button("↑", key=f"up_{i}", disabled=(i == 0)):
                    st.session_state.photos[i], st.session_state.photos[i-1] = \
                        st.session_state.photos[i-1], st.session_state.photos[i]
                    clear_widget_states()
                    st.rerun()
            with o2:
                if st.button("↓", key=f"dn_{i}",
                             disabled=(i == len(st.session_state.photos)-1)):
                    st.session_state.photos[i], st.session_state.photos[i+1] = \
                        st.session_state.photos[i+1], st.session_state.photos[i]
                    clear_widget_states()
                    st.rerun()
            with o3:
                if st.button("🗑", key=f"rm_{i}"):
                    st.session_state.photos.pop(i)
                    clear_widget_states()
                    st.rerun()

        with c2:
            # ── 核心做法：txt_ key 由 Streamlit 管理，AI 生成前先刪 key ──
            txt_key = f"txt_{i}"
            prompt_key = f"prompt_{i}"

            # 初始化：第一次或 AI 更新後，把 photos 的值推入 session_state
            if txt_key not in st.session_state:
                st.session_state[txt_key] = p["text"]
            if prompt_key not in st.session_state:
                st.session_state[prompt_key] = p["prompt"]

            st.text_area(
                label="勘驗文字", height=160,
                placeholder="點擊「✨ AI 生成」自動產製，或直接手動輸入...",
                label_visibility="collapsed",
                key=txt_key
            )
            # 手動編輯同步回 photos
            st.session_state.photos[i]["text"] = st.session_state[txt_key]

            st.text_input(
                label="提示詞",
                placeholder="補充提示詞（選填）",
                label_visibility="collapsed",
                key=prompt_key
            )
            st.session_state.photos[i]["prompt"] = st.session_state[prompt_key]

            if st.button("✨ AI 生成勘驗文字", key=f"gen_{i}"):
                if not st.session_state.api_key:
                    st.error("請先在左側輸入 Gemini API Key")
                else:
                    with st.spinner(f"Gemini AI 分析第 {i+1} 張照片中（含全案脈絡）..."):
                        try:
                            result = call_gemini(
                                st.session_state.api_key,
                                p["bytes"], p["mime"],
                                st.session_state[prompt_key],
                                all_photos=st.session_state.photos,
                                target_index=i,
                                case_num=st.session_state.get("case_num", "")
                            )
                            # 關鍵：先刪 widget key，rerun 後重新初始化才能顯示新值
                            st.session_state.photos[i]["text"] = result
                            del st.session_state[txt_key]
                            st.rerun()
                        except Exception as e:
                            st.error(f"AI 生成失敗：{e}")

        st.markdown("---")

    # ── 批次生成 + 匯出 ───────────────────────────────────
    b1, b2 = st.columns(2)

    with b1:
        if st.button("⚡ 全部照片 AI 生成（含全案脈絡）", use_container_width=True):
            if not st.session_state.api_key:
                st.error("請先輸入 API Key")
            else:
                prog = st.progress(0, text="批次分析中...")
                case_num = st.session_state.get("case_num", "")
                for i in range(len(st.session_state.photos)):
                    p = st.session_state.photos[i]
                    try:
                        result = call_gemini(
                            st.session_state.api_key,
                            p["bytes"], p["mime"],
                            p["prompt"],
                            all_photos=st.session_state.photos,
                            target_index=i,
                            case_num=case_num
                        )
                        st.session_state.photos[i]["text"] = result
                    except Exception as e:
                        st.session_state.photos[i]["text"] = f"（生成失敗：{e}）"
                    prog.progress(
                        (i+1) / len(st.session_state.photos),
                        text=f"處理中 {i+1}/{len(st.session_state.photos)}..."
                    )
                clear_widget_states()
                st.rerun()

    with b2:
        if st.button("📄 產生匯出 docx", use_container_width=True, type="primary"):
            case = st.session_state.get("case_num", "") or "未命名案號"
            doc = Document()
            sec = doc.sections[0]
            sec.page_width = Cm(21); sec.page_height = Cm(29.7)
            sec.left_margin = sec.right_margin = Cm(2.5)
            sec.top_margin = sec.bottom_margin = Cm(2.5)
            doc.styles["Normal"].font.name = "標楷體"
            doc.styles["Normal"].font.size = Pt(12)

            t = doc.add_paragraph()
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = t.add_run("現場照片勘驗報告")
            r.bold = True; r.font.size = Pt(20); r.font.name = "標楷體"

            m = doc.add_paragraph()
            m.alignment = WD_ALIGN_PARAGRAPH.CENTER
            m.add_run(f"案號：{case}　　製作日期：{date.today().strftime('%Y年%m月%d日')}")
            doc.add_paragraph()

            for i, p in enumerate(st.session_state.photos):
                text = p["text"] or "（未填寫勘驗文字）"

                h = doc.add_paragraph()
                hr = h.add_run(f"【編號 {i+1}】{p['name']}")
                hr.bold = True; hr.font.size = Pt(13)
                hr.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
                pPr = h._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bot = OxmlElement("w:bottom")
                bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4")
                bot.set(qn("w:space"), "4"); bot.set(qn("w:color"), "2563EB")
                pBdr.append(bot); pPr.append(pBdr)

                try:
                    ip = doc.add_paragraph()
                    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    ip.add_run().add_picture(io.BytesIO(p["bytes"]), width=Inches(5))
                except Exception:
                    doc.add_paragraph("（照片無法嵌入）")

                for line in text.split("\n"):
                    if line.strip():
                        lp = doc.add_paragraph(line.strip())
                        lp.runs[0].font.name = "標楷體"
                        lp.runs[0].font.size = Pt(12)
                doc.add_paragraph()

            buf = io.BytesIO()
            doc.save(buf); buf.seek(0)
            st.download_button(
                "⬇ 下載 docx 檔案", data=buf,
                file_name=f"{case}-勘驗報告.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            st.success(f"✓ 檔案產製完成：{case}-勘驗報告.docx")
